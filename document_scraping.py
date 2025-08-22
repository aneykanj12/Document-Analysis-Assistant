# I got lazy with the file scraping so these are all Vibe Coded lol
from docx import Document
import re
import textract
from win32com.client import Dispatch
from PyPDF2 import PdfReader

def txt_to_pdf_scraper(filename):
    try:
        with open(filename, "r", encoding="utf-8") as file:
            lines = file.readlines()
            
            cleaned_lines = [line.strip() for line in lines if line.strip()]

            text = " ".join(cleaned_lines)
            
        return text
    except FileNotFoundError:
        return f"Error: {filename} not found."
    
def docx_to_pdf_scraper(filename, add_periods=True):
    try:
        doc = Document(filename)
    except FileNotFoundError:
        return f"Error: {filename} not found."
    except Exception as e:
        return f"Error opening {filename}: {e}"

    chunks = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            chunks.append(t)

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                chunks.append(" ".join(cells))

    cleaned = []
    for t in chunks:
        t = re.sub(r"\s+", " ", t)
        if add_periods and not re.search(r'[.!?]"?$', t):
            t = t + "."
        cleaned.append(t)

    return " ".join(cleaned)


def doc_to_pdf_scraper(filename, add_periods =True):
    try:
        word = Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(filename)
        text = doc.Content.Text
        doc.Close(False)
        word.Quit()
    except Exception as e:
        try:
            word.Quit()
        except:
            pass
        return f"Error opening {filename}: {e}"
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    cleaned = []
    for t in lines:
        t = re.sub(r"\s+", " ", t)
        if add_periods and not re.search(r'[.!?]"?$', t):
            t = t + "."
        cleaned.append(t)
    return " ".join(cleaned)


def pdf_to_pdf_scraper(filename, add_periods=True):
    try:
        reader = PdfReader(filename)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
    except FileNotFoundError:
        return f"Error: {filename} not found."
    except Exception as e:
        return f"Error opening {filename}: {e}"
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    cleaned = []
    for t in lines:
        t = re.sub(r"\s+", " ", t)
        if add_periods and not re.search(r'[.!?]"?$', t):
            t = t + "."
        cleaned.append(t)
    return " ".join(cleaned)

def get_text(file_name):
    file_type = file_name[-3:]

    if(file_type == "txt"):
        return txt_to_pdf_scraper(file_name)
    elif(file_type == "docx"):
        return docx_to_pdf_scraper(file_name)
    elif(file_type == "doc"):
        return doc_to_pdf_scraper(file_name)
    elif(file_type == "pdf"):
        return pdf_to_pdf_scraper(file_name)
    else:
        return "Invalid File Type"